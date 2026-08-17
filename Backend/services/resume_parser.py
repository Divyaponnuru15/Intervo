import os
import fitz
from docx import Document


# =========================================================
# PDF TEXT EXTRACTION
# =========================================================

def extract_text_from_pdf(filepath):

    text = ""

    try:

        if not os.path.exists(filepath):
            print("PDF file does not exist:", filepath)
            return ""

        pdf = fitz.open(filepath)

        for page in pdf:

            page_text = page.get_text("text")

            if page_text:
                text += page_text + "\n"

        pdf.close()

    except Exception as e:

        print("PDF extraction error:", repr(e))

        return ""

    return text.strip()


# =========================================================
# DOCX TEXT EXTRACTION
# =========================================================

def extract_text_from_docx(filepath):

    text_parts = []

    try:

        if not os.path.exists(filepath):
            print("DOCX file does not exist:", filepath)
            return ""

        doc = Document(filepath)

        # Normal paragraphs
        for paragraph in doc.paragraphs:

            if paragraph.text.strip():
                text_parts.append(
                    paragraph.text.strip()
                )

        # Tables
        for table in doc.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:
                        row_text.append(cell_text)

                if row_text:

                    text_parts.append(
                        " | ".join(row_text)
                    )

    except Exception as e:

        print(
            "DOCX extraction error:",
            repr(e)
        )

        return ""

    return "\n".join(text_parts).strip()


# =========================================================
# RESUME TEXT EXTRACTION
# =========================================================

def extract_resume_text(filepath):

    if not filepath:
        return ""

    # IMPORTANT:
    # Do NOT convert the actual filepath to lowercase.
    #
    # Render/Linux filenames are case-sensitive.
    #
    # Only convert the extension for checking.

    extension = os.path.splitext(filepath)[1].lower()


    if extension == ".pdf":

        return extract_text_from_pdf(
            filepath
        )


    elif extension == ".docx":

        return extract_text_from_docx(
            filepath
        )


    else:

        print(
            "Unsupported resume format:",
            extension
        )

        return ""